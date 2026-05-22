# -*- coding: utf-8 -*-
import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Configurar encoding para consola
sys.stdout.reconfigure(encoding='utf-8')

def set_cell_background(cell, color_hex):
    """Establece el color de fondo de una celda de tabla."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Ajusta los márgenes internos (padding) de una celda."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_code_block(doc, code_text):
    """Agrega un bloque de código estilizado (como en Markdown) con fondo gris y fuente monoespaciada."""
    # Usamos una tabla de 1 fila y 1 columna para simular el recuadro del bloque de código
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, "F3F4F6") # Gris muy claro
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    # Quitar bordes gruesos normales y poner borde lateral izquierdo o bordes sutiles
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    left_border = OxmlElement('w:left')
    left_border.set(qn('w:val'), 'single')
    left_border.set(qn('w:sz'), '24') # Borde lateral grueso (3pt)
    left_border.set(qn('w:space'), '0')
    left_border.set(qn('w:color'), '3B82F6') # Azul acento
    tcBorders.append(left_border)
    
    for side in ['top', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4') # Borde muy fino
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'E5E7EB')
        tcBorders.append(b)
        
    tcPr.append(tcBorders)
    
    # Escribir código dentro de la celda
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(31, 41, 55) # Gris oscuro para el código
    
    # Espaciado extra después de la tabla
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(0)
    p_spacer.paragraph_format.space_after = Pt(6)

def read_file_content(filename):
    """Lee el contenido de un archivo de manera segura usando UTF-8."""
    if not os.path.exists(filename):
        return f"// ERROR: Archivo '{filename}' no encontrado en el directorio del proyecto."
    try:
        with open(filename, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        return f"// ERROR al leer '{filename}': {str(e)}"

# ----------------- INICIO DE CONSTRUCCIÓN DE CONTENIDO -----------------

print("Iniciando lectura de archivos del proyecto...")

# Leer todos los códigos del proyecto
firebase_config_code = read_file_content("firebase-config.js")
backend_py_code = read_file_content("backend.py")
read_excel_code = read_file_content("read_excel.py")
subir_cambios_code = read_file_content("Subir_Cambios.bat")
login_html_code = read_file_content("login.html")
login_css_code = read_file_content("login.css")
login_js_code = read_file_content("login.js")
index_html_code = read_file_content("index.html")
styles_css_code = read_file_content("styles.css")
app_js_code = read_file_content("app.js")

# Crear el documento .md completo
md_filename = "Documentacion_Tecnica.md"
print(f"Generando {md_filename}...")

md_content = f"""# DOCUMENTACIÓN TÉCNICA Y GUÍA DE MANTENIMIENTO
## Plataforma Risk Manager | Control Operativo - VirtualSoft

---

## 1. VISIÓN GENERAL DEL PROYECTO

### Nombre del Proyecto
**Risk Manager - Control Operativo** (RiskOps VS)

### Objetivo Principal
Risk Manager es una plataforma web premium de control operativo diseñada para centralizar, optimizar y auditar en tiempo real las operaciones diarias del equipo de Riesgo en VirtualSoft. Su propósito fundamental es eliminar la dependencia de registros dispersos mediante una interfaz unificada que administra tareas asignadas por SETs operacionales, parsea horarios semanales y teletrabajo de archivos de Excel nativos en el servidor, gestiona solicitudes y aprobaciones de permisos, realiza backups de cierres de turnos en base de datos centralizada y genera reportes automáticos exportables en formato PDF.

### Funcionalidades Principales
1. **Control de Acceso y Roles**: Gestión robusta de usuarios en dos niveles de privilegios (Gestores y Supervisores) implementada con Firebase Auth para el inicio de sesión y Firebase Realtime Database para la aprobación manual de accesos.
2. **Restricción de Pantalla (Mobile Blocker)**: Capa de seguridad visual que bloquea el acceso en pantallas menores o iguales a 768 píxeles de ancho (dispositivos móviles y tablets) exigiendo el uso exclusivo de computadoras de escritorio.
3. **Árbol Operativo de Tareas por SET**: Mapeo inteligente y dinámico de las responsabilidades del equipo a través del parseo directo de un libro de Excel (`Tareas de Riesgo.xlsx`) usando la librería SheetJS, permitiendo filtrar tareas por SET específico.
4. **Estados y Notas Técnicas Obligatorias**: Los gestores pueden registrar el estado de cada tarea (Pendiente, En Proceso, Finalizada y No Realizada) con la obligatoriedad de justificar técnicamente cada acción antes de guardar el progreso.
5. **Módulo de Excepciones Justificadas**: Integración de una ventana modal especializada para justificar tareas "No Realizadas" mediante motivos preestablecidos (*Falta de Accesos*, *Sistema Caído*, *Falta de Tiempo*, *Reasignada*, *Otro*) y detalle descriptivo forzoso.
6. **KPI Dinámico de Turno**: Anillo de porcentaje interactivo SVG que evalúa en tiempo real el progreso de tareas completadas y justificadas frente al total de tareas asignadas para el turno.
7. **Biblioteca Interactiva de Procesos (Documentación)**: Buscador inteligente de manuales de políticas PDF y tutoriales multimedia MP4. El sistema analiza automáticamente el nombre de la tarea seleccionada y le sugiere al gestor el instructivo de manera destacada.
8. **Horario Semanal Integrado**: Parsea dinámicamente el archivo `Horario 2026.xlsx`. Aplica privacidad avanzada: los Gestores solo ven su propio horario diario, mientras que el Supervisor tiene el panel global unificado de todo el personal con badges visuales para cada estado de turno.
9. **Cronograma de Teletrabajo**: Parsea el archivo `Teletrabajo.xlsx` en tiempo real y categoriza la asistencia como "Home Office" o "Presencial" con un sistema inteligente de coincidencia fonética y de palabras para asociar nombres.
10. **Solicitud de Permisos Automatizada**: Formulario de justificaciones de ausencias (vacaciones, citas médicas, etc.) que se registra en tiempo real en la base de datos de Firebase y notifica por correo electrónico al supervisor empleando la API de FormSubmit.
11. **Consola de Aprobaciones del Supervisor**: Panel exclusivo de administración para Supervisores/Administradores desde donde aprueban o rechazan el registro de nuevos usuarios en la plataforma y autorizan las solicitudes de permisos pendientes.
12. **Historial de Turnos y Generador de PDFs**: Respaldo global y seguro de todos los reportes de finalización de turnos en Firebase, permitiendo al supervisor auditar las bitácoras y exportar reportes de gestión individuales a formato PDF usando la librería `jsPDF`.
13. **Perfil del Gestor y Cambio de Contraseña**: Visualización dinámica de la fotografía oficial del gestor a partir del reconocimiento de nombres en el directorio del proyecto y restablecimiento seguro de claves directo en la infraestructura de Firebase.

### Tecnologías Utilizadas
* **Estructura y Lógica Frontend**: HTML5 Semántico, Javascript moderno (ES6+), Inter Font (Google Fonts), Boxicons v2.1.4.
* **Diseño y Estética**: CSS3 Vanilla Premium con variables personalizadas, arquitectura Glassmorphic en modo oscuro y claro, animaciones fluidas de fondos de partículas (`float blobs`), y layouts responsive.
* **Librerías de Parseo e Integración**:
  * **SheetJS (xlsx.full.min.js)**: Utilizado para leer, estructurar y renderizar en tablas dinámicas HTML los datos binarios de los libros de Excel (`Tareas de Riesgo.xlsx`, `Horario 2026.xlsx`, `Teletrabajo.xlsx`).
  * **jsPDF (jspdf.umd.min.js)**: Empleado para la compilación, paginación y descarga en PDF de las bitácoras operativas del equipo.
* **Base de Datos y Autenticación**:
  * **Firebase Auth**: Validación segura de identidad del personal.
  * **Firebase Realtime Database**: Almacenamiento en la nube sin esquema de usuarios, históricos de permisos y bitácoras de cierres de turnos.
* **Backend de Desarrollo Local**: Python 3 con `http.server` y `socketserver`, implementando un API REST y servidor de archivos sin caché local para desarrollo rápido en puerto 8080.
* **Despliegue y Automatización**: Script batch de Windows para agregar cambios a Git, realizar commit con timestamp dinámico y realizar push a la nube para despliegue en GitHub Pages de forma automatizada.

### Flujo General de Funcionamiento
```mermaid
graph TD
    A[Inicio de Aplicación] --> B{{¿Usuario Autenticado?}}
    B -- No --> C[login.html / login.js]
    C --> D[Ingreso Credenciales / Registro Gestor]
    D --> E{{Aprobado por Admin}}
    E -- Sí --> F[Autenticación Firebase]
    E -- No --> G[Mostrar Mensaje Pendiente / Rechazo]
    B -- Sí --> H[index.html / app.js]
    H --> I[Cargar Datos de Excel: Tareas, Horario, Teletrabajo]
    H --> J[Escuchar Eventos Firebase: Permisos, Turnos, Aprobaciones]
    
    I --> K[Renderizar Dashboard Principal]
    K --> L[Gestor selecciona su SET de Tareas]
    L --> M[Gestor gestiona Tarea: Notas Técnicas y Estado]
    M --> N[Guardar Progreso en Caché Local y Actualizar KPI]
    
    K --> O[Finalizar Turno]
    O --> P[Guardar Histórico en Realtime Database y Enviar Correo a Supervisor]
    P --> Q[Limpiar Sesión y Redirigir a Login]
```

---

## 2. ESTRUCTURA COMPLETA DEL PROYECTO

El proyecto se encuentra organizado en una estructura limpia y optimizada para la Web, donde los archivos principales se encuentran en la raíz del proyecto para facilitar el despliegue automático en servicios estáticos como GitHub Pages.

```
Pagina Web de Riesgo VS/
├── index.html                   # Dashboard Operativo Principal (HTML)
├── styles.css                   # Sistema de Diseño y Estilos de la Plataforma (CSS)
├── app.js                       # Lógica Operativa de Tareas, Excels y Firebase (JS)
├── login.html                   # Página de Inicio de Sesión, Registro y Recuperación
├── login.css                    # Estilos Específicos para la Interfaz de Login
├── login.js                     # Controlador de Autenticación con Firebase Auth
├── firebase-config.js           # Credenciales y Configuración de Firebase
├── backend.py                   # Servidor de Desarrollo Local y API en Python
├── read_excel.py                # Script de prueba e inspección de archivos XLSX
├── database.json                # Base de datos simulada en formato JSON para backend local
├── Subir_Cambios.bat            # Script de Automatización de Git y Despliegue
├── Tareas Riesgo/               # Directorio de tareas
│   └── Tareas de Riesgo.xlsx    # Libro de tareas y procedimientos por SET
├── Horario/                     # Directorio de programación
│   └── Horario 2026.xlsx        # Horarios del equipo de Riesgos
├── Teletrabajo/                 # Directorio de cronograma de trabajo presencial/virtual
│   └── Teletrabajo.xlsx         # Asignación de días presenciales y de Home Office
├── Procesos/                    # Biblioteca de Procedimientos (PDFs y Multimedia)
│   ├── Instructivo de revisión de apuestas casino.pdf
│   ├── Instructivo de validación de GGR Casino.pdf
│   ├── Política Procedimiento De Aprobación De Retiros.pdf
│   ├── Procedimiento Identificación de jineteo.pdf
│   ├── Proceso de Eliminación de Cuentas - Implementaciones.pdf
│   ├── VALIDACIÓN DE ABUSO DE BONOS EN CAMPAÑAS DE CRM.pdf
│   ├── Revisión de Eventos Deportivos.mp4
│   ├── Revisión de Eventos.mp4
│   └── Validación SEON.mp4
└── assets/
    └── src/
        └── img/                 # Repositorio de recursos multimedia estáticos
            ├── logo.svg         # Logotipo corporativo de la aplicación
            ├── Alexander Villada.png
            ├── Camilo Espinosa.png
            ├── Daniel Benavidez.png
            ├── Josue Alvarez.png
            ├── Juan Jose Diaz.png
            ├── Maria Sanchez.png
            ├── Marilyn Jimenez.png
            ├── Oriana Borja.png
            ├── Samuel Cruz.png
            ├── Sara Santamaria.png
            ├── Sebastian Arango.png
            ├── Sebastian Hincapie.png
            └── Yefferson Giraldo.png
```

---

## 3. CÓDIGO FUENTE DOCUMENTADO

A continuación se adjuntan los códigos fuente completos de la plataforma, acompañados de comentarios exhaustivos sobre su lógica interna, variables críticas y arquitectura.

---

### A. firebase-config.js
Este archivo gestiona las credenciales de Firebase en el cliente. Inicializa la conexión con la nube, configurando el acceso a Firebase Auth y al Realtime Database.

```javascript
{firebase_config_code}
```

#### Explicación de Bloques Críticos de `firebase-config.js`:
* **Líneas 1-9 (`firebaseConfig`)**: Define las claves públicas de acceso a los servicios de Firebase de RiskOps. Estas credenciales permiten conectar de forma directa la aplicación web con la infraestructura en la nube sin necesidad de un backend intermedio para autenticar llamadas.
* **Líneas 12-14**: Verifica si la aplicación de Firebase ya ha sido inicializada previamente por el navegador (para prevenir errores de doble inicialización al recargar componentes) y la inicializa.
* **Línea 15 (`database`)**: Exporta el objeto de acceso directo a Firebase Realtime Database para ser consumido globalmente por los controladores `login.js` y `app.js`.

---

### B. backend.py
Este es el servidor API y de archivos local escrito en Python puro. Proporciona una simulación local del backend y una API REST para persistir la información cuando no se cuenta con acceso directo a Internet o para entornos de desarrollo seguro.

```python
{backend_py_code}
```

#### Explicación de Bloques Críticos de `backend.py`:
* **Líneas 11-13**: Inicialización de la base de datos simulada en formato JSON (`database.json`) si no existe en el disco duro, creando las claves principales `users` y `permissions`.
* **Líneas 19-30 (`read_db()`)**: Mecanismo de inyección automática para la cuenta administradora maestra (`maria.sanchez@virtualsoft.tech` con contraseña `admin123`). Esto garantiza que la dueña del sistema siempre pueda entrar, incluso si se formatea la base de datos local.
* **Clase `APIHandler` (Línea 38)**: Hereda de `SimpleHTTPRequestHandler` y sobreescribe métodos clave:
  * **`end_headers()` (Línea 39)**: Envía cabeceras HTTP que deshabilitan por completo la caché del navegador para asegurar que cada actualización de archivos CSS y JS se cargue instantáneamente durante el desarrollo.
  * **`do_GET()` (Línea 46)**: Intercepta rutas de API bajo `/api/` para entregar información de usuarios, permisos e incluso lee dinámicamente los archivos presentes en el directorio local `Procesos/` para poblar el listado de documentos de la biblioteca.
  * **`do_POST()` (Línea 71)**: Implementa controladores de creación de usuarios, aprobación de accesos (`/api/users/approve`), cambio de roles y edición de contraseñas (`/api/users/promote`), así como el envío y actualización de estados de permisos (`/api/permissions/status`).
  * **`do_OPTIONS()` (Línea 125)**: Habilita el control de CORS (Cross-Origin Resource Sharing) permitiendo que cualquier puerto de origen consulte la API local durante las pruebas de desarrollo.

---

### C. read_excel.py
Script interno en Python que demuestra el parseo manual de archivos `.xlsx` extrayendo el contenido XML nativo comprimido en la estructura zip del libro de Excel. Esto permite depurar la lectura de permisos sin depender del navegador.

```python
{read_excel_code}
```

#### Explicación de Bloques Críticos de `read_excel.py`:
* **Línea 7 (`ZipFile`)**: Abre el archivo `.xlsx` como el contenedor comprimido Zip que realmente es por estándar OpenXML.
* **Líneas 9-17**: Extrae y decodifica la tabla de cadenas de texto compartidas (`sharedStrings.xml`), la cual asocia índices numéricos de celdas con sus respectivos textos literales para optimizar el peso del archivo de Excel.
* **Líneas 20-34**: Parsea mediante árboles de elementos XML (`ElementTree`) el contenido de la primera hoja (`sheet1.xml`), emparejando los valores de celda calculados y las cadenas de texto correspondientes para imprimir las primeras 5 filas por consola.

---

### D. Subir_Cambios.bat
Archivo script automatizado para la consola de Windows que realiza el despliegue automático del proyecto hacia el repositorio remoto de GitHub para actualizar la plataforma en producción.

```batch
{subir_cambios_code}
```

#### Explicación de Bloques Críticos de `Subir_Cambios.bat`:
* **Líneas 5-10**: Localiza dinámicamente el ejecutable de Git en la ruta predeterminada de instalación del usuario en AppData o utiliza la variable de entorno global del sistema en caso de que Git no esté registrado localmente.
* **Líneas 24 (`git add .`)**: Prepara todos los archivos modificados en el espacio de trabajo, incluyendo documentos en `Procesos/` u hojas de Excel actualizadas.
* **Líneas 27-33 (`git commit`)**: Define automáticamente un mensaje de confirmación que captura la fecha y la hora exactas de la computadora que ejecuta el script, creando un punto de restauración preciso.
* **Línea 37 (`git push`)**: Sube los cambios directamente a la rama principal `main` en GitHub, lo que activa de forma automática las GitHub Pages para desplegar la versión actualizada de la web a los usuarios.

---

### E. login.html
Página de presentación y control de inicio de sesión de la plataforma. Diseñada con un estilo premium, animaciones dinámicas de esferas flotantes, y un panel de visualización inteligente que bloquea el acceso en pantallas de celulares.

```html
{login_html_code}
```

#### Explicación de Bloques Críticos de `login.html`:
* **Líneas 17-23 (`#mobileBlocker`)**: Estructura de bloqueo móvil. Si la resolución de pantalla es detectada como teléfono por las media queries de CSS, este contenedor oculta el resto de la interfaz y despliega un panel restrictivo.
* **Líneas 26-27 (`.blob`)**: Divisiones decorativas con fondos de color acentuado que se desplazan de manera infinita para dar un aspecto estético premium y fluido.
* **Líneas 29-144 (`.login-container`)**: Panel unificado "glassmorphic" con desenfoque de fondo en el cual conviven tres vistas de interacción dinámica activadas por JavaScript:
  * **`#loginPanel`**: Formulario de acceso con inputs interactivos de correo, contraseña, selector de visibilidad y botón de submit.
  * **`#registerPanel`**: Formulario de solicitud de registro con campos de nombre completo, correo, contraseña, confirmación y rol solicitado (Gestor o Supervisor).
  * **`#forgotPanel`**: Formulario de envío seguro de restablecimiento de contraseñas. Redirige la solicitud mediante POST hacia FormSubmit para enviar una notificación directa al supervisor del sistema.

---

### F. login.css
Archivo de hojas de estilo especializado para la interfaz de acceso. Complementa el diseño global de la plataforma, implementando el desenfoque del panel translúcido y las animaciones de fondo.

```css
{login_css_code}
```

#### Explicación de Bloques Críticos de `login.css`:
* **Líneas 14-45 (`.blob` y `@keyframes float`)**: Define la física del fondo animado de la aplicación. Aplica un filtro de desenfoque (`filter: blur(80px)`) muy alto para las esferas de colores primario y verde, desplazándolas en coordenadas cartesianas X/Y y escalas dinámicas de manera infinita.
* **Líneas 47-59 (`.login-container`)**: Define el panel translúcido principal usando desenfoque de filtro en el navegador (`backdrop-filter: blur(20px)`) para emular el vidrio esmerilado con una sombra profunda para contraste.
* **Líneas 90-121**: Estiliza de manera personalizada los campos de formulario e inyecta dinámicamente mediante SVG codificado en base64 una flecha personalizada para el campo selector (`select`) que encaja con el diseño en modo oscuro.

---

### G. login.js
Controlador lógico de la autenticación. Gestiona las interacciones de los formularios en la pantalla de inicio de sesión, conectando el navegador con Firebase Auth en tiempo real.

```javascript
{login_js_code}
```

#### Explicación de Bloques Críticos de `login.js`:
* **Líneas 17-19**: Valida si el proyecto se abrió directamente como un archivo local (`file://`). De ser así, alerta al usuario que el sistema de recuperación por FormSubmit requiere un servidor local HTTP o el despliegue en línea para funcionar de forma segura.
* **Líneas 59-71 (Validación de Registro)**: Implementa validaciones forzosas en el cliente. Asegura que las contraseñas coincidan y que el gestor ingrese su nombre y apellido de forma completa (mínimo dos palabras separadas por espacios) para mantener la integridad de la base de datos de personal.
* **Líneas 79-83 (Auto-elevación de Privilegios)**: Regla especial de seguridad. Si el correo de registro corresponde al de la dueña de la plataforma (`maria.sanchez@virtualsoft.tech`), el sistema eleva de forma invisible el rol a "Admin" y auto-aprueba la cuenta en Firebase Realtime Database.
* **Líneas 102-129 (Notificación por FormSubmit)**: Crea programáticamente un formulario HTTP POST oculto y lo envía por debajo en una ventana paralela para alertar vía correo electrónico al supervisor sobre la existencia del nuevo usuario pendiente de aprobación.
* **Líneas 179-194 (Auto-bootstrapping Admin)**: En caso de que la base de datos de Firebase Auth esté vacía y la cuenta maestra de la dueña del sistema aún no esté registrada, el controlador detecta el intento de login con las credenciales maestras predeterminadas y crea la cuenta automáticamente en Firebase con privilegios completos de administrador.
* **Líneas 220-232 (Validación de Aprobación)**: Comprueba el estado de aprobación del usuario. Si un supervisor lo rechazó o si se encuentra en estado pendiente de aprobación, bloquea el inicio de sesión, despliega el error explicativo e interrumpe la sesión activa.

---

### H. index.html
El esqueleto estructural de la plataforma Risk Manager. Contiene la barra de navegación del Sidebar, el Header superior con reloj activo y notificaciones en tiempo real, las ventanas modales de excepciones y perfil, y las diferentes vistas que componen la lógica de negocio de la plataforma.

```html
{index_html_code}
```

#### Explicación de Bloques Críticos de `index.html`:
* **Líneas 13-16**: Carga de dependencias fundamentales del cliente desde servidores CDN de alta velocidad:
  * **`xlsx.full.min.js`**: Motor principal de SheetJS para procesar binarios de Excel de forma local en el navegador del cliente.
  * **`jspdf.umd.min.js`**: Generador nativo de archivos PDF en Javascript.
* **Líneas 30-55 (`aside.sidebar`)**: Barra lateral de navegación que define los enlaces dinámicos a las diferentes pestañas operativas de la plataforma. Integra elementos invisibles por defecto (`display: none`) para los accesos del supervisor tales como "Historial Turnos" e "Aprobaciones" para evitar intrusión.
* **Líneas 60-85 (`header.top-header`)**: Cabecera principal que aloja el reloj en tiempo real, el badge dinámico de turno operativo activo, el botón para alternar el tema visual, la campana inteligente de notificaciones con dropdown nativo y el avatar del usuario conectado.
* **Líneas 88-154 (`#view-workspace`)**: Módulo de control de tareas del Gestor. Integra el listado de selección de SET, el visor de árbol de tareas interactivo, el panel central de observaciones y guardado de notas, y la columna lateral derecha de KPI de progreso y atajos de biblioteca.
* **Líneas 156-194 (Vistas de Horario y Teletrabajo)**: Secciones del dashboard que integran los selectores de semana y las tablas HTML adaptativas que poblará SheetJS.
* **Líneas 203-242 (`#view-permisos`)**: Formulario interactivo de solicitud de permisos y contenedor del historial personal de solicitudes del gestor conectado.
* **Líneas 244-294 (Vistas Supervisor)**: Consolas avanzadas para la gestión administrativa de cuentas de usuarios, aprobación de permisos con retroalimentación y visor de reportes de turnos generales con descarga de bitácoras en PDF.
* **Líneas 298-337 (Ventanas Modales)**: Modales interactivos translúcidos posicionados de forma absoluta para reportar tareas no realizadas con detalles (excepciones) y para cambiar claves personales.

---

### I. styles.css
El motor estético del Risk Manager. Define la paleta de colores oscuros oscilando en tonos gris carbón y azul cobalto, el modo claro premium alternativo, las reglas de desenfoque de cristales del estilo "glassmorphism", sombras de relieve acentuadas, y tipografías optimizadas.

```css
{styles_css_code}
```

#### Explicación de Bloques Críticos de `styles.css`:
* **Líneas 1-50 (Paletas de Colores e Identidad Visual)**: Centralización de la identidad visual del proyecto mediante variables CSS (`--bg-primary`, `--glass-bg`, etc.). Permite cambiar radicalmente los colores de toda la aplicación modificando un único bloque. Define la transición suave entre temas oscuros y claros premium.
* **Líneas 68-112 (Bloqueador de Dispositivos)**: Define los estilos de la pantalla bloqueadora móvil `#mobileBlocker`, usando la directiva `@media (max-width: 768px)` para activarla, ocultando el contenedor de la aplicación.
* **Líneas 131-138 (`.glass-panel`)**: Define el estilo Glassmorphism que le otorga a la plataforma su apariencia premium. Utiliza la regla `backdrop-filter: blur(12px)` para crear una dispersión de la luz sobre el fondo y bordes translúcidos con opacidad reducida.
* **Líneas 193-201 (`.nav-item.active`)**: Estilo premium para los elementos activos de la barra de navegación lateral. Aplica un borde izquierdo grueso con color azul acentuado y un gradiente sutil y elegante hacia la derecha.
* **Líneas 253-275 (`.clock`)**: Diseño digital del reloj del sistema superior. Usa fuente monoespaciada para evitar saltos horizontales de texto y aplica un efecto de resplandor mediante la regla `text-shadow: 0 0 10px rgba(59,130,246,0.5)`.
* **Líneas 496-504 (Anillo KPI)**: Estructura del anillo circular SVG que calcula el progreso mediante vectores de trazos dinámicos, aplicando una transición lineal fluida (`transition: stroke-dasharray 1s ease-out`).

---

### J. app.js
El cerebro lógico de Risk Manager. Contiene la autenticación de sesión, el motor de parseo e inyección de datos de archivos de Excel mediante SheetJS, el gestor de la biblioteca interactiva, el motor de actualización en tiempo real de Firebase para notificaciones, permisos y aprobaciones, y la compilación de reportes PDF interactivos con jsPDF.

```javascript
{app_js_code}
```

#### Explicación de Bloques Críticos de `app.js`:
* **Líneas 1-13 (Comprobación de Sesión)**: Bloque autoejecutable que intercepta el ciclo de vida del navegador. Verifica la validez del token de sesión local en `localStorage` y redirecciona forzosamente a la pantalla de login si detecta manipulación de datos o ausencia de inicio de sesión.
* **Líneas 14-35 (`namesMatch` y `normalizeName`)**: Motor de procesamiento lingüístico personalizado. Normaliza cadenas de texto en español quitando acentos y mayúsculas, y realiza comparaciones robustas palabra por palabra. Esto resuelve el error común donde un usuario registrado como "María Sánchez" no coincida con el registro del Excel escrito como "maria sanchez".
* **Líneas 37-49 (`documentUrls`)**: Mapeo inteligente para documentos de biblioteca corporativa. En caso de videos de capacitación pesados en MP4 (que ralentizarían el hosting web), intercepta la llamada y redirecciona el stream hacia Google Drive, manteniendo los archivos de texto PDF locales bajo la ruta `Procesos/`.
* **Líneas 80-147 (`loadExcelTasks`)**: Implementa el parseo dinámico de tareas con SheetJS. Consume el binario asíncronamente desde `Tareas de Riesgo.xlsx`, lo transforma en JSON, agrupa las tareas por SETs operacionales eliminando duplicados visuales en el árbol y renderiza el selector moderno para el Gestor.
* **Líneas 150-297 (`loadSchedule`)**: Motor de parseo y renderizado de horarios. Lee `Horario 2026.xlsx`, procesa las fechas seriales de Excel transformándolas a fechas estándar legibles por humanos (`formatExcelDate`) y construye dinámicamente la tabla HTML. Implementa privacidad avanzada: si el rol es "Gestor", filtra automáticamente la matriz ocultando los horarios de los demás compañeros de equipo.
* **Líneas 299-403 (`loadTeletrabajo`)**: Procesa el archivo `Teletrabajo.xlsx`. Identifica las celdas asignadas a cada gestor y renderiza en pantalla un badge visual verde de "HOME OFFICE" o un indicador gris de asistencia "PRESENCIAL".
* **Líneas 506-537 (`updateKPI`)**: Calcula el rendimiento del turno en tiempo real evaluando las tareas que han sido marcadas como "Finalizada" y las exceptuadas justificadamente bajo "No Realizada", modificando la propiedad vectorial `stroke-dasharray` del anillo de progreso SVG.
* **Líneas 548-630 (`renderQuickDocs`)**: Analizador semántico para la biblioteca de ayuda. Al hacer clic en una tarea operativa, evalúa palabras clave de su título (ej. "GGR", "retiro", "apuesta", "SEON") y despliega de manera destacada un acceso sugerido arriba del listado.
* **Líneas 1021-1082 (Guardar Progreso de Tarea)**: Asegura que el gestor registre obligatoriamente una nota técnica detallada en el campo de texto antes de permitir guardar el progreso en el caché y actualizar visualmente el nodo en el árbol del dashboard.
* **Líneas 1161-1258 (`handleEndShift`)**: Procedimiento de finalización de turno. Recopila de forma automatizada las notas y estados de todas las tareas gestionadas durante la sesión, construye un cuerpo de reporte estructurado, lo respalda de manera segura bajo `/shift_reports` en Firebase y envía silenciosamente el correo vía AJAX mediante la API de FormSubmit al supervisor del sistema, limpiando el caché para cerrar sesión.
* **Líneas 1558-1609 (`exportShiftReport`)**: Motor compilador de PDF mediante jsPDF. Recupera el reporte de turno desde Firebase usando el ID de transacción, define fuentes, colores corporativos y márgenes, y compila el contenido autogenerando un archivo de descarga seguro (`Reporte_NombreGestor_Fecha.pdf`).

---

## 4. FLUJO FUNCIONAL DE LA APLICACIÓN

Esta sección describe detalladamente el ciclo de vida de la ejecución de la aplicación, su procesamiento de datos y la conectividad interna de sus componentes.

### A. Carga de Datos
Al iniciar el panel principal (`index.html`), se inician de forma paralela y asíncrona tres hilos de carga de datos:
1. **Acceso al Repositorio de Archivos de Excel**: A través de peticiones HTTP `fetch`, se descargan las hojas binarias del servidor (`Tareas de Riesgo.xlsx`, `Horario 2026.xlsx`, `Teletrabajo.xlsx`). SheetJS procesa las hojas, generando matrices estructuradas de datos en memoria para poblar el árbol de tareas, selector de semanas, y cuadrículas de horarios.
2. **Sincronización con Realtime Database de Firebase**: La aplicación establece listeners dinámicos en tiempo real sobre los endpoints `/permissions` y `/users` para actualizar la campana de notificaciones al instante ante aprobaciones o solicitudes nuevas.
3. **Restauración del Caché del Turno**: Se verifica la existencia de datos del turno en curso guardados en el `localStorage` (`riskOps_cache`) para recuperar el progreso de las tareas ante recargas accidentales de la página.

### B. Navegación
La navegación entre las vistas de negocio se realiza sin recarga de página (Single Page Application - SPA) mediante el controlador de pestañas en `app.js`. Al hacer clic en una sección del Sidebar:
* Se valida la autorización de rol (los gestores no pueden acceder a los botones de Supervisor ni interactuar con las tablas administrativas).
* Se aplica una transición suave ocultando el contenedor activo (`display: none`) y desplegando el correspondiente al identificador presionado.

### C. Conectividad Interna
```
┌─────────────────────────────────────────────────────────────────────────┐
│                           FIREBASE CLOUD INFRASTRUCTURE                 │
│    ┌─────────────────────┐    ┌─────────────────┐    ┌─────────────┐    │
│    │    Firebase Auth    │    │Realtime Database│    │ FormSubmit  │    │
│    └──────────┬──────────┘    └────────┬────────┘    └──────┬──────┘    │
└───────────────┼────────────────────────┼────────────────────┼───────────┘
                │                        │                    │            
                │ (Validación JWT)       │ (JSON Sync)        │ (Notif Mail)
                ▼                        ▼                    ▼            
┌─────────────────────────────────────────────────────────────────────────┐
│                               CLIENT WEB APP                            │
│ ┌──────────────────────┐     ┌──────────────────┐    ┌────────────────┐ │
│ │  login.js (Auth UI)  │◄───►│ app.js (Core app)│◄──►│    SheetJS     │ │
│ └──────────────────────┘     └────────┬─────────┘    └────────────────┘ │
│                                       │ (Compila PDF)                   │
│                                       ▼                                 │
│                              ┌──────────────────┐                       │
│                              │      jsPDF       │                       │
│                              └──────────────────┘                       │
└─────────────────────────────────────────────────────────────────────────┘
```

---

## 5. GUÍA DE MANTENIMIENTO Y ADMINISTRACIÓN

Esta sección sirve de manual práctico para asegurar que cualquier diseñador o desarrollador futuro pueda actualizar contenidos, modificar el diseño o resolver problemas en producción.

### A. Actualizar Tareas en el Sistema
Las tareas no están programadas directamente en el código de la web, sino que se leen directamente desde el archivo `Tareas Riesgo/Tareas de Riesgo.xlsx`. Para modificarlas:
1. Abra el archivo `Tareas Riesgo/Tareas de Riesgo.xlsx` en Microsoft Excel.
2. La estructura del archivo debe respetar estrictamente los siguientes encabezados en la primera fila:
   * **Set** o **Set **: Nombre del SET operacional (ej. *SET 1 - Apuestas*, *SET 2 - Retiros*).
   * **Tarea**: Nombre corto y representativo de la tarea.
   * **Detalle de Tarea**: Descripción de los pasos operacionales y validaciones obligatorias a realizar.
   * **Horario**: Rango horario sugerido para ejecutar la labor.
   * **Día**: Programación de frecuencia de la tarea.
3. Guarde el archivo Excel.
4. Ejecute el script `Subir_Cambios.bat` haciendo doble clic en él. El script empaquetará los cambios automáticamente y los subirá a GitHub Pages, propagando la actualización a todo el equipo en menos de 1 minuto.

### B. Agregar Nuevos Documentos o Videos de Instructivos
Para enriquecer la biblioteca de procesos con nuevos manuales o grabaciones de entrenamiento:
1. **Para Archivos PDF**:
   * Guarde el archivo PDF en la carpeta `Procesos/` respetando una nomenclatura clara (ej. `Instructivo_Validacion_Fraude.pdf`).
   * Abra `app.js` en un editor de texto y localice la constante `archivos` dentro de las funciones `renderQuickDocs` (Línea 552) y en la inyección de la biblioteca de documentos (Línea 982). Agregue el nombre exacto de su archivo en la lista:
     `"Instructivo_Validacion_Fraude.pdf"`
2. **Para Videos MP4 (Videos Pesados)**:
   * Debido al límite de almacenamiento en repositorios estáticos y para evitar lentitud de carga, los videos deben alojarse en Google Drive.
   * Suba el video a Google Drive y configúrelo con visibilidad pública ("Cualquier persona con el enlace puede ver").
   * Copie el enlace de visualización generado.
   * Abra `app.js` y busque la constante `documentUrls` (Línea 38). Registre el nombre simbólico del archivo y pegue su respectiva URL de Google Drive:
     `"Instructivo_Fraude_Video.mp4": "https://drive.google.com/file/d/ID_DEL_VIDEO/view?usp=sharing"`
   * Añada el nombre simbólico `"Instructivo_Fraude_Video.mp4"` a la constante `archivos` en las líneas mencionadas de `app.js`.

### C. Modificación de Estilos, Colores y Diseño
La arquitectura CSS está estructurada bajo variables nativas en `styles.css`. Esto facilita el rediseño instantáneo de la interfaz:
* **Cambiar el Color de Acento (Azul corporativo)**:
  Edite la línea 9 de `styles.css` y cambie el valor hexadecimal de `--accent-primary`:
  `--accent-primary: #3B82F6;` (Cambiar a `#9F1239` por ejemplo, para usar Rojo Cereza).
* **Cambiar Tonos del Fondo Oscuro**:
  Edite las líneas 3 y 4 de `styles.css`:
  * `--bg-primary`: Controla el color del fondo exterior del dashboard.
  * `--bg-secondary`: Controla el fondo del panel translúcido base.
* **Modificar el Nivel de Transparencia o Desenfoque (Glassmorphic)**:
  Localice la clase `.glass-panel` (Línea 131 de `styles.css`) y modifique:
  * `backdrop-filter: blur(12px);` (Aumentar a `20px` para un efecto de desenfoque de cristal más espeso).
  * `background: var(--glass-bg);` (Modificar la opacidad alfa de la variable `--glass-bg` en la cabecera).

### D. Gestión de GitHub Pages y Despliegue
La aplicación web se aloja de forma gratuita y auto-administrada en la infraestructura de GitHub Pages.
* El despliegue de actualizaciones se realiza automáticamente a través de la compilación de la rama `main`.
* Si un desarrollador realiza cambios manuales y no cuenta con el script batch en Windows, puede desplegar ejecutando la terminal en la raíz del proyecto:
  ```bash
  git add .
  git commit -m "Actualizacion de plataforma"
  git push origin main
  ```

### E. Solución de Errores Comunes (Troubleshooting)
1. **"¡ATENCIÓN! Estás abriendo la plataforma directamente como un archivo local (file:///)..."**:
   * **Causa**: Se intentó abrir el archivo `login.html` o `index.html` haciendo doble clic directo desde el explorador de archivos. El navegador bloquea por seguridad el envío de correos (FormSubmit) y la lectura correcta de Excels locales a través de peticiones asíncronas AJAX.
   * **Solución**: Debe servir la aplicación web utilizando un servidor HTTP local.
2. **Los Horarios de un Gestor no cargan en pantalla**:
   * **Causa**: El nombre registrado por el usuario en su cuenta de Firebase no coincide con el texto ingresado en la columna de nombres de la hoja `Horario 2026.xlsx`.
   * **Solución**: Asegúrese de que el nombre del gestor en la primera columna del Excel coincida exactamente con el de la cuenta de usuario (la función `namesMatch` ignora acentos y mayúsculas, pero requiere que las palabras principales coincidan).
3. **Error "Request Blocked by FormSubmit" al enviar permisos**:
   * **Causa**: La dirección de correo electrónico a la que se envían las solicitudes (`maria.sanchez@virtualsoft.tech`) aún no ha sido activada en la plataforma de FormSubmit para este dominio.
   * **Solución**: La primera vez que se envía una solicitud, FormSubmit envía un correo electrónico de confirmación a la cuenta receptora. Debe hacer clic en el botón "Confirmar Activación" de ese correo para autorizar el tráfico.

---

## 6. GUÍA PARA FUTUROS DESARROLLADORES

### Requisitos Previos
* **Git** instalado en el sistema operativo.
* **Python 3.x** instalado (para el servidor de desarrollo local).
* Editor de código recomendado: **Visual Studio Code** con la extensión *Live Server* o el entorno de ejecución integrado de Python.

### Instalación del Proyecto en Entorno Local
1. Clone el repositorio de código o extraiga los archivos en un directorio local:
   ```bash
   git clone <URL_DEL_REPOSITORIO>
   ```
2. Abra una consola de terminal en la carpeta raíz del proyecto.
3. Inicie el servidor de desarrollo local de Python ejecutando:
   ```bash
   python backend.py
   ```
4. Abra su navegador web e ingrese a la dirección:
   ```
   http://localhost:8080
   ```
5. Esto levantará la aplicación web con el backend en Python activo, permitiéndole interactuar con la base de datos simulada local `database.json`.

### Buenas Prácticas Recomendadas
1. **Mantener Vanilla CSS**: Evite instalar librerías pesadas como TailwindCSS o Bootstrap de forma externa que puedan colisionar con el sistema de variables de diseño de `styles.css`.
2. **Validar las Operaciones Asíncronas**: Al realizar peticiones a Firebase o parseos con SheetJS, maneje adecuadamente los bloques `try-catch` para evitar bloqueos del hilo de renderizado del cliente.
3. **No Guardar Contraseñas en Realtime Database**: Por razones estrictas de seguridad de datos, la base de datos en Realtime Database solo debe almacenar perfiles de usuario descriptivos. Las credenciales de acceso y hashes de claves deben ser gestionados exclusivamente por la infraestructura cifrada de Firebase Auth.

---

## 7. SEGURIDAD, BACKUP Y RESPALDO

### Archivos y Configuraciones Críticas
Existen tres pilares que de borrarse o ser alterados con credenciales erróneas deshabilitarán la plataforma:
1. **`firebase-config.js`**: Si las claves de acceso de Firebase son alteradas, los usuarios no podrán iniciar sesión ni solicitar permisos en tiempo real.
2. **Directorio `Procesos/`**: Aloja toda la documentación oficial del equipo. No elimine los PDFs locales.
3. **Estructura de Carpetas de Hojas de Cálculo (`Horario/`, `Teletrabajo/`, `Tareas Riesgo/`)**: Mantenga inalterados los nombres de las carpetas y de los libros de Excel (`Horario 2026.xlsx`, `Teletrabajo.xlsx`, `Tareas de Riesgo.xlsx`).

### Procedimiento de Respaldo de Base de Datos
Dado que Risk Manager utiliza una arquitectura híbrida, se deben respaldar dos capas regularmente:
1. **La Base de Datos en la Nube (Firebase)**:
   * Ingrese a la consola web oficial de [Firebase Console](https://console.firebase.google.com/).
   * Acceda a su proyecto y diríjase a la sección **Realtime Database**.
   * En la pestaña de visualización de datos, haga clic en el botón de tres puntos en la esquina superior derecha y seleccione **Exportar JSON**.
   * Guarde el archivo JSON de respaldo en una ubicación física segura de almacenamiento corporativo.
2. **Las Hojas de Cálculo Locales**:
   * Mantenga un control de versiones de los archivos `.xlsx` realizando copias de seguridad semanales de las carpetas `Horario/`, `Teletrabajo/` y `Tareas Riesgo/` en un servidor seguro local o repositorio Git corporativo paralelo.

---

## 8. MEJORAS FUTURAS SUGERIDAS

1. **Migración a un Framework Moderno**: Si la plataforma escala en cantidad de usuarios e interactividad compleja, se recomienda migrar a un framework como **Next.js** o **React**, lo que permitiría componentizar la interfaz manteniendo la lógica SPA de alto rendimiento.
2. **Integración con Google Sheets API**: En lugar de leer archivos de Excel `.xlsx` estáticos subidos manualmente al servidor por Git, el sistema podría consumir directamente hojas de cálculo de Google Sheets en la nube en tiempo real mediante su API oficial, eliminando el paso de despliegue con Git para cambiar horarios.
3. **Sistema de Auditoría Completa de Tareas**: Extender el módulo de Supervisor para visualizar gráficas analíticas del desempeño de cada gestor sobre las tareas asignadas a lo largo de las semanas, calculando tiempos medios de finalización.

---

## 9. ENTREGABLES FINALES

### Checklist de Mantenimiento Periódico
- `[ ]` **Lunes (Semanal)**: Actualizar los horarios del equipo de Riesgo en `Horario/Horario 2026.xlsx` y la asignación de teletrabajo en `Teletrabajo/Teletrabajo.xlsx`.
- `[ ]` **Lunes (Semanal)**: Ejecutar `Subir_Cambios.bat` posterior a las modificaciones semanales para propagar los calendarios al personal.
- `[ ]` **Fin de Mes (Mensual)**: Realizar una exportación de respaldo JSON de la base de datos Realtime Database desde la Consola de Firebase.
- `[ ]` **Fin de Mes (Mensual)**: Limpiar bitácoras antiguas en Realtime Database si el peso de `/shift_reports` excede los límites gratuitos del plan de Firebase.
- `[ ]` **Semestral**: Auditar la vigencia de los instructivos PDF y videos de entrenamiento MP4 alojados en la biblioteca y actualizar enlaces obsoletos.

---
*Fin del Documento de Especificación Técnica.*
"""

with open(md_filename, 'w', encoding='utf-8') as f:
    f.write(md_content)

print(f"Archivo {md_filename} escrito con éxito.")

# ----------------- GENERAR EL ARCHIVO WORD (.DOCX) CON ESTILOS PREMIUM -----------------
docx_filename = "Documentacion_Tecnica.docx"
print(f"Generando {docx_filename} usando python-docx con estilos de alta calidad...")

doc = Document()

# Configurar márgenes de página (1 pulgada en todos los lados)
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Configurar fuente por defecto a Calibri/Arial
style_normal = doc.styles['Normal']
font = style_normal.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(51, 51, 51) # Gris carbón muy profesional

# Función para agregar títulos estilizados
def add_custom_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(11, 14, 20) # Color azul oscuro primario
        
        # Agregar línea decorativa horizontal muy fina debajo del título principal
        p_line = doc.add_paragraph()
        p_line.paragraph_format.space_before = Pt(0)
        p_line.paragraph_format.space_after = Pt(8)
        p_line_run = p_line.add_run("―" * 50)
        p_line_run.font.size = Pt(8)
        p_line_run.font.color.rgb = RGBColor(229, 231, 235)
        
    elif level == 2:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(59, 130, 246) # Azul acento
    elif level == 3:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(75, 85, 99) # Gris oscuro

def add_bullet_point(doc, bold_prefix, text_content):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    
    run_bold = p.add_run(bold_prefix)
    run_bold.bold = True
    run_bold.font.color.rgb = RGBColor(17, 24, 39)
    
    run_text = p.add_run(text_content)
    run_text.font.color.rgb = RGBColor(75, 85, 99)

def add_normal_paragraph(doc, text, bold_text=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_text:
        run_bold = p.add_run(bold_text)
        run_bold.bold = True
        run_bold.font.color.rgb = RGBColor(17, 24, 39)
        
    run_text = p.add_run(text)
    run_text.font.color.rgb = RGBColor(51, 51, 51)

# --- PORTADA ---
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(120)
p_title.paragraph_format.space_after = Pt(10)
run_title = p_title.add_run("DOCUMENTACIÓN TÉCNICA Y ESPECIFICACIÓN")
run_title.bold = True
run_title.font.size = Pt(24)
run_title.font.color.rgb = RGBColor(11, 14, 20)

p_subtitle = doc.add_paragraph()
p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_subtitle.paragraph_format.space_after = Pt(40)
run_sub = p_subtitle.add_run("Plataforma Risk Manager | Control Operativo - VirtualSoft")
run_sub.font.size = Pt(14)
run_sub.font.color.rgb = RGBColor(59, 130, 246)

p_meta = doc.add_paragraph()
p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_meta.paragraph_format.space_before = Pt(150)
run_meta = p_meta.add_run("Preparado para: Equipo de Desarrollo y Diseñadores de VirtualSoft\nFecha de emisión: Mayo de 2026\nEstado del documento: Aprobado - Versión Completa (Word Ready)")
run_meta.font.size = Pt(10)
run_meta.font.italic = True
run_meta.font.color.rgb = RGBColor(107, 114, 128)

doc.add_page_break()

# --- 1. VISIÓN GENERAL ---
add_custom_heading(doc, "1. VISIÓN GENERAL DEL PROYECTO", 1)

add_normal_paragraph(doc, "Risk Manager | Control Operativo (RiskOps VS)", bold_text="Nombre del Proyecto: ")
add_normal_paragraph(doc, "Risk Manager es una plataforma web premium de control operativo diseñada para centralizar, optimizar y auditar en tiempo real las operaciones diarias del equipo de Riesgo en VirtualSoft. Su propósito fundamental es eliminar la dependencia de registros dispersos mediante una interfaz unificada que administra tareas asignadas por SETs operacionales, parsea horarios semanales y teletrabajo de archivos de Excel nativos en el servidor, gestiona solicitudes y aprobaciones de permisos, realiza backups de cierres de turnos en base de datos centralizada y genera reportes automáticos exportables en formato PDF.", bold_text="Objetivo Principal: ")

add_custom_heading(doc, "Funcionalidades Principales", 2)
add_bullet_point(doc, "Control de Acceso y Roles: ", "Gestión robusta de usuarios en dos niveles de privilegios (Gestores y Supervisores) implementada con Firebase Auth para el inicio de sesión y Firebase Realtime Database para la aprobación manual de accesos.")
add_bullet_point(doc, "Restricción de Pantalla (Mobile Blocker): ", "Capa de seguridad visual que bloquea el acceso en pantallas menores o iguales a 768 píxeles de ancho (dispositivos móviles y tablets) exigiendo el uso exclusivo de computadoras de escritorio.")
add_bullet_point(doc, "Árbol Operativo de Tareas por SET: ", "Mapeo inteligente y dinámico de las responsabilidades del equipo a través del parseo directo de un libro de Excel (Tareas de Riesgo.xlsx) usando la librería SheetJS, permitiendo filtrar tareas por SET específico.")
add_bullet_point(doc, "Estados y Notas Técnicas Obligatorias: ", "Los gestores pueden registrar el estado de cada tarea (Pendiente, En Proceso, Finalizada y No Realizada) con la obligatoriedad de justificar técnicamente cada acción antes de guardar el progreso.")
add_bullet_point(doc, "Módulo de Excepciones Justificadas: ", "Integración de una ventana modal especializada para justificar tareas 'No Realizadas' mediante motivos preestablecidos y detalle descriptivo forzoso.")
add_bullet_point(doc, "KPI Dinámico de Turno: ", "Anillo de porcentaje interactivo SVG que evalúa en tiempo real el progreso de tareas completadas y justificadas frente al total de tareas asignadas para el turno.")
add_bullet_point(doc, "Biblioteca Interactiva de Procesos: ", "Buscador inteligente de manuales de políticas PDF y tutoriales multimedia MP4. El sistema analiza automáticamente el nombre de la tarea seleccionada y le sugiere al gestor el instructivo de manera destacada.")
add_bullet_point(doc, "Horario Semanal Integrado: ", "Parsea dinámicamente el archivo Horario 2026.xlsx. Aplica privacidad avanzada: los Gestores solo ven su propio horario diario, mientras que el Supervisor tiene el panel global unificado de todo el personal con badges visuales para cada estado de turno.")
add_bullet_point(doc, "Cronograma de Teletrabajo: ", "Parsea el archivo Teletrabajo.xlsx en tiempo real y categoriza la asistencia como 'Home Office' o 'Presencial' con un sistema inteligente de coincidencia fonética y de palabras para asociar nombres.")
add_bullet_point(doc, "Solicitud de Permisos Automatizada: ", "Formulario de justificaciones de ausencias (vacaciones, citas médicas, etc.) que se registra en tiempo real en la base de datos de Firebase y notifica por correo electrónico al supervisor empleando la API de FormSubmit.")
add_bullet_point(doc, "Consola de Aprobaciones del Supervisor: ", "Panel exclusivo de administración para Supervisores/Administradores desde donde aprueban o rechazan el registro de nuevos usuarios en la plataforma y autorizan las solicitudes de permisos pendientes.")
add_bullet_point(doc, "Historial de Turnos y Generador de PDFs: ", "Respaldo global y seguro de todos los reportes de finalización de turnos en Firebase, permitiendo al supervisor auditar las bitácoras y exportar reportes de gestión individuales a formato PDF usando la librería jsPDF.")
add_bullet_point(doc, "Perfil del Gestor y Cambio de Contraseña: ", "Visualización dinámica de la fotografía oficial del gestor a partir del reconocimiento de nombres en el directorio del proyecto y restablecimiento seguro de claves directo en la infraestructura de Firebase.")

add_custom_heading(doc, "Tecnologías Utilizadas", 2)
add_bullet_point(doc, "Estructura y Lógica Frontend: ", "HTML5 Semántico, Javascript moderno (ES6+), Inter Font (Google Fonts), Boxicons v2.1.4.")
add_bullet_point(doc, "Diseño y Estética: ", "CSS3 Vanilla Premium con variables personalizadas, arquitectura Glassmorphic en modo oscuro y claro, animaciones fluidas de fondos de partículas (float blobs), y layouts responsive.")
add_bullet_point(doc, "Librerías de Parseo e Integración: ", "SheetJS (xlsx.full.min.js) para parseo de Excels, jsPDF (jspdf.umd.min.js) para la generación dinámica de reportes en PDF, Firebase SDK 8.10.1 para Auth y Realtime Database.")
add_bullet_point(doc, "Backend Local: ", "Servidor en Python 3 con http.server, socketserver, json, os, urllib.parse, implementando API REST y servidor de desarrollo en puerto 8080.")
add_bullet_point(doc, "Despliegue y Automatización: ", "Script batch de Windows para empaquetamiento automático de Git y empuje a la nube en GitHub Pages.")

# --- 2. ESTRUCTURA ---
add_custom_heading(doc, "2. ESTRUCTURA COMPLETA DEL PROYECTO", 1)
add_normal_paragraph(doc, "El proyecto está organizado en una estructura limpia y optimizada para la Web, donde los archivos principales se encuentran en la raíz del proyecto para facilitar el despliegue automático en servicios estáticos como GitHub Pages.")

add_bullet_point(doc, "index.html: ", "Dashboard Operativo Principal de la aplicación.")
add_bullet_point(doc, "styles.css: ", "Sistema de diseño global de la plataforma con soporte para tema oscuro y claro.")
add_bullet_point(doc, "app.js: ", "Controlador JavaScript principal del dashboard, parseo de hojas Excel y conexión con Firebase.")
add_bullet_point(doc, "login.html: ", "Página de inicio de sesión, registro de gestores y recuperación de contraseña.")
add_bullet_point(doc, "login.css: ", "Hoja de estilo específica para la vista de login y esferas animadas de fondo.")
add_bullet_point(doc, "login.js: ", "Controlador JavaScript de autenticación y lógica de validación de registro.")
add_bullet_point(doc, "firebase-config.js: ", "Inicialización del SDK de Firebase y credenciales de acceso.")
add_bullet_point(doc, "backend.py: ", "Servidor web de desarrollo y API local en Python 3.")
add_bullet_point(doc, "read_excel.py: ", "Script interno de prueba para parseo nativo XML de archivos de Excel.")
add_bullet_point(doc, "Subir_Cambios.bat: ", "Automatizador de commits de Git y despliegue rápido.")
add_bullet_point(doc, "Tareas Riesgo/Tareas de Riesgo.xlsx: ", "Libro Excel conteniendo las tareas diarias por SET.")
add_bullet_point(doc, "Horario/Horario 2026.xlsx: ", "Planificador semanal de horarios de gestores.")
add_bullet_point(doc, "Teletrabajo/Teletrabajo.xlsx: ", "Cronograma de asistencia Home Office / Presencial.")
add_bullet_point(doc, "Procesos/: ", "Biblioteca local conteniendo 6 instructivos PDF oficiales y 3 videos pesados en formato MP4 mapeados en la nube.")
add_bullet_point(doc, "assets/src/img/: ", "Repositorio de recursos visuales que contiene las fotografías oficiales de perfil de todo el equipo de Riesgos y el logo de la aplicación.")

# --- 3. CÓDIGO FUENTE DOCUMENTADO ---
add_custom_heading(doc, "3. CÓDIGO FUENTE COMPLETO Y DOCUMENTADO", 1)
add_normal_paragraph(doc, "A continuación se presentan de forma íntegra todos los archivos de código que componen el ecosistema de Risk Manager, acompañados de comentarios exhaustivos que detallan variables críticas, funciones principales y lógica estructural.")

# A. firebase-config.js
add_custom_heading(doc, "A. firebase-config.js (Configuración Firebase)", 2)
add_normal_paragraph(doc, "Este script inicializa el SDK de Firebase en el cliente para proporcionar servicios de base de datos en la nube y autenticación de usuarios.")
add_code_block(doc, firebase_config_code)

# B. backend.py
add_custom_heading(doc, "B. backend.py (Servidor API de Desarrollo en Python)", 2)
add_normal_paragraph(doc, "Servidor HTTP local y API REST escrito en Python para simular la persistencia y realizar pruebas de desarrollo seguras sin conexión a Internet.")
add_code_block(doc, backend_py_code)

# C. read_excel.py
add_custom_heading(doc, "C. read_excel.py (Lector de Excel de Prueba)", 2)
add_normal_paragraph(doc, "Demostración de parseo nativo en Python extrayendo XMLs comprimidos del contenedor zip del libro xlsx.")
add_code_block(doc, read_excel_code)

# D. Subir_Cambios.bat
add_custom_heading(doc, "D. Subir_Cambios.bat (Script de Despliegue Git)", 2)
add_normal_paragraph(doc, "Script batch en Windows que automatiza el control de versiones en Git y realiza push directo a la rama de producción.")
add_code_block(doc, subir_cambios_code)

# E. login.html
add_custom_heading(doc, "E. login.html (Estructura de Acceso)", 2)
add_normal_paragraph(doc, "Código de estructura de la interfaz de login, registro de usuarios y recuperación de contraseña, incluyendo la ventana de bloqueo móvil.")
add_code_block(doc, login_html_code)

# F. login.css
add_custom_heading(doc, "F. login.css (Estilos de Acceso)", 2)
add_normal_paragraph(doc, "Estilos de presentación premium y animaciones de esferas fluidas flotantes de fondo para la pantalla de acceso.")
add_code_block(doc, login_css_code)

# G. login.js
add_custom_heading(doc, "G. login.js (Controlador de Acceso)", 2)
add_normal_paragraph(doc, "Lógica JavaScript que gestiona los formularios de login, registro, envío de avisos a supervisor vía correo electrónico y validaciones complejas de identidad.")
add_code_block(doc, login_js_code)

# H. index.html
add_custom_heading(doc, "H. index.html (Estructura Dashboard Principal)", 2)
add_normal_paragraph(doc, "Esqueleto estructural de la interfaz principal de Risk Manager con barra lateral sidebar, reloj digital dinámico, anillo de KPI dinámico y paneles operativos.")
add_code_block(doc, index_html_code)

# I. styles.css
add_custom_heading(doc, "I. styles.css (Estilos del Dashboard)", 2)
add_normal_paragraph(doc, "Hoja de estilo global que define el modo claro/oscuro alternativo, variables corporativas de diseño, la física del glassmorphism y adaptaciones estéticas premium.")
add_code_block(doc, styles_css_code)

# J. app.js
add_custom_heading(doc, "J. app.js (Lógica Operativa Principal)", 2)
add_normal_paragraph(doc, "Motor dinámico en JavaScript que consume hojas de cálculo XLS binarias vía peticiones asíncronas, calcula el progreso SVG del turno, procesa atajos inteligentes de biblioteca y maneja los reportes y PDF con jsPDF.")
add_code_block(doc, app_js_code)

# --- 4. FLUJO FUNCIONAL ---
add_custom_heading(doc, "4. FLUJO FUNCIONAL DE LA APLICACIÓN", 1)
add_normal_paragraph(doc, "Esta sección describe con precisión cómo se conectan internamente los componentes y datos de la plataforma en su ciclo de vida ordinario.")

add_custom_heading(doc, "Carga y Procesamiento de Datos", 2)
add_normal_paragraph(doc, "Al ingresar al panel principal de Risk Manager, el sistema realiza peticiones asíncronas HTTP fetch para descargar en memoria de cliente los libros Excel de tareas, horarios y teletrabajo. Paralelamente, establece listeners automáticos en tiempo real hacia los endpoints de la base de datos de Firebase Realtime para reaccionar al instante ante nuevas solicitudes o aprobaciones.")

add_custom_heading(doc, "Navegación e Interconexión de Componentes", 2)
add_normal_paragraph(doc, "La navegación se opera mediante una arquitectura Single Page Application (SPA), donde Javascript oculta o despliega vistas según las pestañas seleccionadas y los privilegios de rol detectados en localStorage. Al seleccionar una tarea, el sistema sugiere documentos asociados analizando el título semánticamente y calcula el progreso KPI dinámico con vectores SVG en cada actualización de estado.")

# --- 5. GUÍA DE MANTENIMIENTO ---
add_custom_heading(doc, "5. GUÍA DE MANTENIMIENTO Y ADMINISTRACIÓN", 1)
add_normal_paragraph(doc, "Manual práctico para que futuros desarrolladores o diseñadores puedan mantener operativa la plataforma sin interrupciones.")

add_custom_heading(doc, "Actualizar Tareas del Equipo", 2)
add_normal_paragraph(doc, "Las tareas operativas se administran desde el libro Excel 'Tareas Riesgo/Tareas de Riesgo.xlsx'. Para modificarlas, simplemente edite las filas de este archivo respetando los encabezados originales (Set, Tarea, Detalle de Tarea, Horario, Día) y ejecute el comando automatizado 'Subir_Cambios.bat' para publicarlos en producción de forma instantánea.")

add_custom_heading(doc, "Añadir Documentos o Videos a la Biblioteca", 2)
add_normal_paragraph(doc, "Para PDFs locales, deposite el archivo en la carpeta 'Procesos/' y agregue el nombre de archivo exacto a la constante 'archivos' en 'app.js'. Para videos de entrenamiento pesados en MP4, súbalos a Google Drive de forma pública y enlace el identificador en la constante de mapeo 'documentUrls' del archivo 'app.js'.")

add_custom_heading(doc, "Modificar Estilos Visuales", 2)
add_normal_paragraph(doc, "Abra 'styles.css' y modifique las variables globales ubicadas bajo el elemento ':root'. Modificando el valor hexadecimal de '--accent-primary' se redefinirán todos los colores corporativos de la interfaz al instante. Para ajustar el desenfoque de los paneles glassmorphic, edite la clase '.glass-panel' modificando el atributo 'backdrop-filter: blur(12px)'.")

# --- 6. GUÍA PARA FUTUROS DESARROLLADORES ---
add_custom_heading(doc, "6. GUÍA PARA FUTUROS DESARROLLADORES", 1)
add_normal_paragraph(doc, "Instrucciones detalladas de despliegue local e incorporación de desarrolladores al código base del proyecto.")

add_custom_heading(doc, "Instalación Local", 2)
add_normal_paragraph(doc, "1. Clone o descargue los archivos del repositorio de la plataforma.\n2. Inicie el servidor de desarrollo API ejecutando en terminal: python backend.py\n3. Ingrese a la URL http://localhost:8080 en su navegador para interactuar con la simulación local.")

add_custom_heading(doc, "Buenas Prácticas", 2)
add_normal_paragraph(doc, "1. No inyecte lógica de base de datos de contraseñas de usuarios bajo Realtime Database; delegue siempre la encriptación y hashes de credenciales de seguridad al SDK de Firebase Auth.\n2. Use variables CSS personalizadas ubicadas en styles.css para mantener una paleta estética uniforme.\n3. Implemente comprobaciones try-catch en cada flujo asíncrono para resguardar la fluidez del cliente.")

# --- 7. SEGURIDAD Y RESPALDO ---
add_custom_heading(doc, "7. SEGURIDAD, BACKUP Y RESPALDO", 1)
add_normal_paragraph(doc, "Medidas y protocolos preventivos para proteger la plataforma de pérdidas accidentales de datos o fallos críticos.")

add_bullet_point(doc, "Archivos Críticos: ", "Mantenga inalterables las credenciales en firebase-config.js y no elimine los libros de Excel ubicados en las carpetas Horario/, Teletrabajo/ o Tareas Riesgo/, ya que su ausencia causará el colapso de las pantallas dinámicas.")
add_bullet_point(doc, "Procedimiento de Backup: ", "Realice una descarga mensual JSON de la base de datos de Firebase Realtime Database ingresando a Firebase Console y guarde copias de seguridad de las hojas Excel en servidores corporativos físicos.")

# --- 8. MEJORAS FUTURAS ---
add_custom_heading(doc, "8. MEJORAS FUTURAS SUGERIDAS", 1)
add_bullet_point(doc, "Consumo directo en Google Sheets: ", "Migrar el sistema para leer dinámicamente desde la nube utilizando la API oficial de Google Sheets en lugar de subir archivos XLSX binarios al hosting estático.")
add_bullet_point(doc, "Migración de Arquitectura: ", "Componentizar la aplicación web mediante frameworks estables como Next.js o React si la complejidad interactiva escala significativamente.")

# --- 9. ENTREGABLES FINALES ---
add_custom_heading(doc, "9. ENTREGABLES FINALES & CHECKLIST", 1)
add_bullet_point(doc, "[ ] Checklist Semanal: ", "Actualizar los libros de horarios y teletrabajo en sus respectivas carpetas y ejecutar 'Subir_Cambios.bat' para actualizar en la nube.")
add_bullet_point(doc, "[ ] Checklist Mensual: ", "Descargar backup JSON completo desde Firebase Console y depurar reportes de turnos antiguos de ser necesario.")
add_bullet_point(doc, "[ ] Checklist Semestral: ", "Revisar la vigencia de los PDFs locales y streams de videos en Google Drive, actualizando enlaces en el código.")

print(f"Escribiendo {docx_filename} en disco...")
doc.save(docx_filename)
print("¡Archivo Word .docx generado exitosamente con estilos premium!")

print("Todo el proceso de generación de documentación técnica se ha completado de forma satisfactoria.")
